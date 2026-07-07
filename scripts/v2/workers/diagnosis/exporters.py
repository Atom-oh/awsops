"""AWSops v2 — AI Diagnosis report exporters.

Render an already-generated report markdown into download formats:
  - to_docx(markdown) -> bytes   (python-docx; pure-python)
  - to_pdf(markdown)  -> bytes   (markdown→HTML→headless chromium via playwright)

Read-only: these only transcode a report that was already produced over redacted data. No AWS
mutation, no network egress (the PDF CSS uses the system Noto CJK font — no external @import).
"""
import io
import re

_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^\s*[-*]\s+(.*)$")
_LINK = re.compile(r"\[([^\]]+)\]\([^)]*\)")
_ITALIC_LINE = re.compile(r"^_(.+)_$")
_ORDERED = re.compile(r"^\s*\d+[.)]\s+(.*)$")

# Korean-native sans, ships with every Windows Word install. Deliberately NOT the same font as
# the PDF path's Noto Sans CJK KR: a PDF renders server-side (where Noto CJK is installed in the
# worker image), but a DOCX resolves fonts on the READER's machine, where Noto CJK usually is not
# installed and would fall back to the exact inconsistent-font mess this is fixing.
_FONT = "Malgun Gothic"
_MONO = "Consolas"
_INK_900, _INK_800, _INK_500 = "14130F", "1F1E1D", "5F5A4D"
_BRAND, _BRAND_DEEP, _BRAND_TINT = "D97757", "8E4830", "F5DCCF"
_BORDER, _CODE_BG = "D7D3C7", "F7F6F2"


def _set_east_asia(style_or_run, font):
    """python-docx's font.name setter only writes w:ascii/w:hAnsi — Korean text still renders in
    Word's East-Asian FALLBACK font unless w:eastAsia is also set. No high-level API for this."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    rpr = style_or_run.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def _styled_document():
    """Document() with page setup + base typography explicit — never rely on Word's defaults
    (US Letter, default Calibri, default blue Heading styles)."""
    from docx import Document
    from docx.shared import Mm, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)  # A4, matches the PDF path
    section.top_margin = section.bottom_margin = Mm(18)
    section.left_margin = section.right_margin = Mm(16)

    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(_INK_800)
    _set_east_asia(normal, _FONT)

    # Override Word's default heading-blue: titles stay near-black for readability (color is an
    # ACCENT, never the title text itself); Heading 3 — the level the LLM actually authors via its
    # own `###` — gets the deep-brand tint so AI-authored subsections read as visually distinct
    # from our own H1/H2 report scaffold.
    for name, size, color in (
        ("Heading 1", Pt(20), _INK_900),
        ("Heading 2", Pt(15), _INK_900),
        ("Heading 3", Pt(12.5), _BRAND_DEEP),
    ):
        style = doc.styles[name]
        style.font.name = _FONT
        style.font.size = size
        style.font.color.rgb = RGBColor.from_string(color)
        _set_east_asia(style, _FONT)

    _add_bottom_rule(doc.styles["Heading 1"], _BRAND)

    return doc


def _add_bottom_rule(style, color):
    """A colored bottom border under a heading style — no high-level python-docx API for
    paragraph borders (w:pBdr)."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    ppr = style.element.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def _inline(text: str) -> str:
    """Strip the inline markdown the report uses (links→text, bold/italic/code markers)."""
    text = _LINK.sub(r"\1", text)
    text = text.replace("**", "").replace("`", "")
    return text.strip()


_INLINE_SPLIT = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def _add_runs(paragraph, text: str):
    """Split `**bold**`/`` `code` `` into real runs instead of stripping the markers to plain
    text — the only two sites this corpus actually uses (the deterministic '**목차**' TOC label
    and the coverage note's `` `key` `` spans). Links are stripped first (TOC-only; a downloaded
    ops report gains nothing from clickable anchors)."""
    text = _LINK.sub(r"\1", text).strip()
    for segment in _INLINE_SPLIT.split(text):
        if not segment:
            continue
        if segment.startswith("**") and segment.endswith("**"):
            run = paragraph.add_run(segment[2:-2])
            run.bold = True
        elif segment.startswith("`") and segment.endswith("`"):
            from docx.shared import Pt, RGBColor

            run = paragraph.add_run(segment[1:-1])
            run.font.name = _MONO
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(_BRAND_DEEP)
            _set_east_asia(run, _MONO)
        else:
            paragraph.add_run(segment)
    return paragraph


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def _is_table_separator(line: str) -> bool:
    # |---|:--:| style row — only dashes/colons/spaces/pipes AND at least one dash (so a data row
    # like `|  |  |` with no dashes is NOT mistaken for a separator and dropped).
    return bool(re.fullmatch(r"\s*\|[\s|:-]*-[\s|:-]*\|\s*", line))


def _row_cells(line: str):
    return [_inline(c) for c in line.strip().strip("|").split("|")]


def _shd(pr_element, fill):
    """Inject <w:shd w:val="clear" w:fill="..."/> into a *Pr element (tcPr or pPr) — cell/
    paragraph shading has no high-level python-docx API."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pr_element.append(shd)


def _calm_borders(table):
    """A uniform light-grey grid instead of Word's default loud accent-color table style — no
    high-level API for w:tblBorders."""
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), _BORDER)
        borders.append(el)
    table._tbl.tblPr.append(borders)


def to_docx(markdown: str) -> bytes:
    """Pragmatic markdown→DOCX: headings, the date blockquote, bullets, tables, paragraphs."""
    doc = _styled_document()
    lines = markdown.splitlines()
    i = 0
    seen_h2 = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        # fenced code block — the LLM occasionally emits stray ``` fences (report.py's own
        # _balance_code_fences admits as much); render verbatim as monospace, never through
        # _inline/_add_runs (code content must not be mangled by the bold/code-span parser).
        if stripped.startswith("```"):
            i += 1  # consume the opening fence
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):  # consume the closing fence if one was found (EOF otherwise)
                i += 1
            for code_line in code_lines:
                p = doc.add_paragraph()
                run = p.add_run(code_line)
                run.font.name = _MONO
                _set_east_asia(run, _MONO)
                _shd(p._p.get_or_add_pPr(), _CODE_BG)
            continue

        # table block (consecutive | … | rows)
        if _is_table_row(line):
            raw = []
            while i < len(lines) and _is_table_row(lines[i]):
                raw.append(lines[i])
                i += 1
            has_header = len(raw) >= 2 and _is_table_separator(raw[1])
            block = [_row_cells(r) for r in raw if not _is_table_separator(r)]
            if block:
                cols = max(len(r) for r in block)
                table = doc.add_table(rows=0, cols=cols)
                table.style = "Table Grid"
                table.autofit = False
                _calm_borders(table)
                from docx.shared import Mm

                col_width = Mm(178) // cols
                for c in range(cols):
                    table.columns[c].width = col_width
                for row_idx, cells in enumerate(block):
                    row = table.add_row().cells
                    for c in range(cols):
                        cell = row[c]
                        cell.width = col_width
                        text = cells[c] if c < len(cells) else ""
                        run = cell.paragraphs[0].add_run(text)
                        if has_header and row_idx == 0:
                            run.bold = True
                            _shd(cell._tc.get_or_add_tcPr(), _BRAND_TINT)
            continue

        m = _HEADING.match(stripped)
        if m:
            level = min(len(m.group(1)), 4)
            heading = doc.add_heading(_inline(m.group(2)), level=level)
            if level == 2:
                # Fresh page per major (report-scaffold) section, except the first — the PDF
                # sibling gets this implicitly from print pagination; docx needs it explicit.
                if seen_h2:
                    heading.paragraph_format.page_break_before = True
                seen_h2 = True
            i += 1
            continue

        b = _BULLET.match(line)
        if b:
            _add_runs(doc.add_paragraph(style="List Bullet"), b.group(1))
            i += 1
            continue

        o = _ORDERED.match(line)
        if o:
            # ponytail: built-in "List Number" numbering continues across separate lists
            # document-wide; this corpus has exactly one numbered list (SPOF, sections.py:73) —
            # add numbering-restart oxml only if a second numbered list ever actually appears.
            _add_runs(doc.add_paragraph(style="List Number"), o.group(1))
            i += 1
            continue

        if stripped.startswith(">"):  # blockquote (e.g. the generation-date line) → italic
            from docx.shared import RGBColor

            p = doc.add_paragraph()
            run = p.add_run(_inline(stripped.lstrip("> ").rstrip()))
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(_INK_500)
            i += 1
            continue

        u = _ITALIC_LINE.match(stripped)
        if u:  # whole-line `_..._` (e.g. the degraded-section fallback body) → italic, muted
            from docx.shared import RGBColor

            run = doc.add_paragraph().add_run(_inline(u.group(1)))
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(_INK_500)
            i += 1
            continue

        _add_runs(doc.add_paragraph(), stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# A4 print CSS. System font only (Noto Sans CJK KR is installed in the worker image) — NO external
# @import / Google-Fonts URL, since the Fargate worker runs in a private subnet (an external font
# fetch would hang until timeout and Korean would render as tofu anyway).
_PDF_CSS = """
@page { size: A4; margin: 18mm 16mm; }
* { font-family: 'Noto Sans CJK KR', 'Noto Sans', sans-serif; }
body { font-size: 11pt; line-height: 1.55; color: #1a1a1a; }
h1 { font-size: 20pt; } h2 { font-size: 15pt; margin-top: 1.2em; } h3 { font-size: 12.5pt; }
blockquote { color: #555; border-left: 3px solid #ccc; margin: 0 0 1em; padding: .2em .8em; }
table { border-collapse: collapse; width: 100%; margin: .6em 0; }
th, td { border: 1px solid #bbb; padding: 5px 8px; font-size: 10pt; text-align: left; }
code, pre { font-family: monospace; background: #f4f4f4; }
"""


def _html(md_text: str) -> str:
    import markdown as _md

    body = _md.markdown(md_text, extensions=["tables", "fenced_code"])
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<style>{_PDF_CSS}</style></head><body>{body}</body></html>"
    )


def to_pdf(md_text: str) -> bytes:
    """markdown→HTML→headless chromium→PDF bytes. playwright is imported lazily (image-only dep)."""
    from playwright.sync_api import sync_playwright

    html = _html(md_text)
    with sync_playwright() as p:
        # Fargate blocks the user-namespace sandbox chromium uses by default → --no-sandbox; and the
        # container runs unprivileged so the setuid sandbox helper can't be used either →
        # --disable-setuid-sandbox. Safe here: JS is disabled below and the HTML is static, server-
        # built from already-redacted report data (no untrusted scripts to contain).
        browser = p.chromium.launch(headless=True, args=["--no-sandbox", "--disable-setuid-sandbox"])
        try:
            # Static render of LLM-generated HTML over redacted data → no script execution.
            page = browser.new_page(java_script_enabled=False)
            page.set_content(html, wait_until="load")
            return page.pdf(format="A4", print_background=True)
        finally:
            browser.close()
