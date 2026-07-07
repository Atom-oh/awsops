import io

from diagnosis import exporters

_SAMPLE = "# AWS 진단 리포트\n\n> 생성 일시: 2026-06-17 09:00 (KST)\n\n## 요약\n\n본문 문단입니다.\n\n- 항목 A\n- 항목 B\n\n| 키 | 값 |\n|----|----|\n| a  | 1  |\n"


def _doc(md):
    from docx import Document

    return Document(io.BytesIO(exporters.to_docx(md)))


def test_to_docx_returns_docx_zip_bytes():
    out = exporters.to_docx(_SAMPLE)
    assert isinstance(out, (bytes, bytearray)) and len(out) > 0
    assert bytes(out[:4]) == b"PK\x03\x04"  # DOCX is a zip (OOXML)


def test_to_docx_preserves_content():
    import io
    from docx import Document

    doc = Document(io.BytesIO(exporters.to_docx(_SAMPLE)))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "AWS 진단 리포트" in text   # heading rendered
    assert "본문 문단입니다." in text   # body paragraph rendered
    assert any(t.rows for t in doc.tables)  # the markdown table became a docx table


def test_to_pdf_returns_pdf_bytes():
    import pytest
    pytest.importorskip("playwright")  # playwright is image-only; skip in local/CI without it
    try:
        out = exporters.to_pdf(_SAMPLE)
    except Exception as e:  # chromium binary not installed in this env
        pytest.skip(f"chromium unavailable: {e}")
    assert isinstance(out, (bytes, bytearray)) and bytes(out[:5]) == b"%PDF-"


def test_html_template_uses_system_font_no_external_import():
    html = exporters._html("# t\n\n본문")
    assert "Noto Sans CJK KR" in html
    assert "@import" not in html and "fonts.googleapis" not in html  # no egress in private subnet


def test_docx_page_setup_a4_explicit_margins():
    # Round to whole mm: OOXML stores page size in integer twips, so Mm(210) round-trips as
    # ~210.0086mm (the standard A4-in-twips rounding), not an exact EMU match.
    sec = _doc(_SAMPLE).sections[0]
    assert round(sec.page_width.mm) == 210
    assert round(sec.page_height.mm) == 297
    assert round(sec.top_margin.mm) == 18
    assert round(sec.bottom_margin.mm) == 18
    assert round(sec.left_margin.mm) == 16
    assert round(sec.right_margin.mm) == 16


def test_docx_normal_style_korean_font_and_east_asia():
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    st = _doc(_SAMPLE).styles["Normal"]
    assert st.font.name == "Malgun Gothic"
    assert st.font.size == Pt(10.5)
    assert st.font.color.rgb == RGBColor.from_string("1F1E1D")
    rfonts = st.element.rPr.find(qn("w:rFonts"))
    assert rfonts is not None
    assert rfonts.get(qn("w:eastAsia")) == "Malgun Gothic"


def test_docx_heading_styles_not_word_default():
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    styles = _doc(_SAMPLE).styles
    expect = [
        ("Heading 1", Pt(20), "14130F"),
        ("Heading 2", Pt(15), "14130F"),
        ("Heading 3", Pt(12.5), "8E4830"),  # brand-deep — distinguishes LLM ### subheads
    ]
    for name, size, color in expect:
        st = styles[name]
        assert st.font.size == size, name
        assert st.font.color.rgb == RGBColor.from_string(color), name
        rfonts = st.element.rPr.find(qn("w:rFonts"))
        assert rfonts is not None and rfonts.get(qn("w:eastAsia")) == "Malgun Gothic", name


def test_docx_h1_has_brand_bottom_rule():
    from docx.oxml.ns import qn

    st = _doc(_SAMPLE).styles["Heading 1"]
    pbdr = st.element.pPr.find(qn("w:pBdr"))
    assert pbdr is not None
    bottom = pbdr.find(qn("w:bottom"))
    assert bottom is not None
    assert bottom.get(qn("w:val")) == "single"
    assert bottom.get(qn("w:color")) == "D97757"


def test_docx_page_break_before_second_h2_only():
    md = "# t\n\n## A\n\nbody\n\n## B\n\nbody\n"
    doc = _doc(md)
    h2s = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert len(h2s) == 2
    assert not h2s[0].paragraph_format.page_break_before
    assert h2s[1].paragraph_format.page_break_before is True


def test_docx_table_explicit_widths_everywhere():
    from docx.shared import Mm

    t = _doc(_SAMPLE).tables[0]
    assert t.autofit is False
    # Round to whole twips-of-a-mm: OOXML stores widths in integer twips, so an exact EMU
    # equality fails by the same rounding as the page-size check above.
    w_mm = round((Mm(178) // len(t.columns)) / 36000)  # 1mm == 36000 EMU
    for col in t.columns:
        assert round(col.width.mm) == w_mm
    for row in t.rows:
        for cell in row.cells:
            assert round(cell.width.mm) == w_mm


def test_docx_table_header_bold_and_shaded():
    from docx.oxml.ns import qn

    t = _doc(_SAMPLE).tables[0]  # `| 키 | 값 |` header + separator + one data row
    header = t.rows[0]
    for cell in header.cells:
        assert cell.paragraphs[0].runs and all(r.bold for r in cell.paragraphs[0].runs)
        shd = cell._tc.tcPr.find(qn("w:shd"))
        assert shd is not None
        assert shd.get(qn("w:fill")) == "F5DCCF"
    data_row = t.rows[1]
    for cell in data_row.cells:
        shd = cell._tc.tcPr.find(qn("w:shd")) if cell._tc.tcPr is not None else None
        assert shd is None


def test_docx_table_no_header_gets_no_bold_or_shading():
    from docx.oxml.ns import qn

    md = "| a | b |\n| c | d |\n"  # two data rows, no `|---|---|` separator → headerless
    t = _doc(md).tables[0]
    for row in t.rows:
        for cell in row.cells:
            assert not any(r.bold for r in cell.paragraphs[0].runs)
            shd = cell._tc.tcPr.find(qn("w:shd")) if cell._tc.tcPr is not None else None
            assert shd is None


def test_docx_table_calm_border_color():
    from docx.oxml.ns import qn

    t = _doc(_SAMPLE).tables[0]
    assert t.style.name == "Table Grid"
    top = t._tbl.tblPr.find(qn("w:tblBorders")).find(qn("w:top"))
    assert top.get(qn("w:color")) == "D7D3C7"
    assert top.get(qn("w:sz")) == "4"


def test_docx_toc_label_real_bold_no_literal_asterisks():
    md = "# t\n\n**목차**\n\n- [항목](#a)\n"
    doc = _doc(md)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "**" not in full_text
    toc_label = next(p for p in doc.paragraphs if p.text == "목차")
    assert toc_label.runs and all(r.bold for r in toc_label.runs)


def test_docx_inline_code_monospace_run():
    md = "# t\n\n- `inventory`: ok\n"
    doc = _doc(md)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "`" not in full_text
    item = next(p for p in doc.paragraphs if p.text == "inventory: ok")
    code_run = next(r for r in item.runs if r.text == "inventory")
    assert code_run.font.name == "Consolas"
    plain_run = next(r for r in item.runs if r.text == ": ok")
    assert plain_run.font.name != "Consolas"


def test_docx_underscore_line_italic_no_literal_underscores():
    from docx.shared import RGBColor

    md = "# t\n\n_이 섹션 생성에 실패했습니다 (degraded): boom_\n"
    doc = _doc(md)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "_" not in full_text
    p = next(p for p in doc.paragraphs if "degraded" in p.text)
    assert p.runs and all(r.italic for r in p.runs)
    assert p.runs[0].font.color.rgb == RGBColor.from_string("5F5A4D")


def test_docx_blockquote_muted_color():
    from docx.shared import RGBColor

    p = next(p for p in _doc(_SAMPLE).paragraphs if "생성 일시" in p.text)
    assert p.runs[0].italic is True
    assert p.runs[0].font.color.rgb == RGBColor.from_string("5F5A4D")
